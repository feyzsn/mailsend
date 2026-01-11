import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import pandas as pd
from docx import Document
from docx2pdf import convert
import win32com.client as win32
import os
import pythoncom
import ttkbootstrap as ttk # Modern arayüz kütüphanesi
from ttkbootstrap.constants import * 
class MailGondericiUygulamasi:
    def __init__(self, root):
        self.root = root
        self.root.title("Toplu Mail Asistanı")
        self.root.geometry("700x800")
        
        # Pencere ikonu (varsa)
        try:
            self.root.iconbitmap("logo.ico") 
        except:
            pass

        # Değişkenler
        self.word_path = tk.StringVar()
        self.excel_path = tk.StringVar()
        self.konu_basligi = tk.StringVar()
        self.dosya_eki = tk.StringVar()
        
        self.arayuz_olustur()

    def arayuz_olustur(self):
        # Ana çerçeve
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=BOTH, expand=YES)

        # Başlık
        lbl_baslik = ttk.Label(main_frame, text="Mail Asistanı", font=("Helvetica", 18, "bold"))
        lbl_baslik.pack(pady=(0, 20))

        # --- 1. DOSYA SEÇİMLERİ ---
        # HATA DÜZELTİLDİ: padding parantez içinden silindi.
        file_frame = ttk.LabelFrame(main_frame, text=" Dosya Seçimleri ")
        # Padding buraya (pack içine) taşındı: ipadx ve ipady iç boşluk sağlar
        file_frame.pack(fill=X, pady=5, ipadx=10, ipady=10)

        # Word
        ttk.Label(file_frame, text="Word Şablonu:").pack(anchor=W, padx=5)
        frm_word = ttk.Frame(file_frame)
        frm_word.pack(fill=X, pady=(0, 10), padx=5)
        ttk.Entry(frm_word, textvariable=self.word_path).pack(side=LEFT, fill=X, expand=YES, padx=(0, 5))
        ttk.Button(frm_word, text="Seç", command=self.word_sec, style="secondary.Outline.TButton").pack(side=RIGHT)

        # Excel
        ttk.Label(file_frame, text="Excel Listesi:").pack(anchor=W, padx=5)
        frm_excel = ttk.Frame(file_frame)
        frm_excel.pack(fill=X, padx=5)
        ttk.Entry(frm_excel, textvariable=self.excel_path).pack(side=LEFT, fill=X, expand=YES, padx=(0, 5))
        ttk.Button(frm_excel, text="Seç", command=self.excel_sec, style="secondary.Outline.TButton").pack(side=RIGHT)

        # --- 2. MAIL AYARLARI ---
        # HATA DÜZELTİLDİ: padding parantez içinden silindi.
        mail_frame = ttk.LabelFrame(main_frame, text=" Mail Detayları ")
        # Padding pack içine eklendi
        mail_frame.pack(fill=BOTH, expand=YES, pady=15, ipadx=10, ipady=10)

        # Dosya Eki İsmi
        ttk.Label(mail_frame, text="PDF Dosya Ek Adı (Opsiyonel):").pack(anchor=W, padx=5)
        ttk.Entry(mail_frame, textvariable=self.dosya_eki).pack(fill=X, pady=(0, 5), padx=5)
        ttk.Label(mail_frame, text="* Boş bırakırsanız sadece isim kullanılır.", font=("Helvetica", 8)).pack(anchor=W, pady=(0, 10), padx=5)

        # Konu
        ttk.Label(mail_frame, text="Mail Konusu:").pack(anchor=W, padx=5)
        ttk.Entry(mail_frame, textvariable=self.konu_basligi).pack(fill=X, pady=(0, 10), padx=5)

        # Gövde
        ttk.Label(mail_frame, text="Mail İçeriği:").pack(anchor=W, padx=5)
        self.text_area_body = ttk.Text(mail_frame, height=5)
        self.text_area_body.pack(fill=BOTH, expand=YES, pady=5, padx=5)

        # --- 3. AKSİYON VE LOG ---
        # Başlat Butonu
        ttk.Button(main_frame, text="GÖNDERİMİ BAŞLAT", command=self.islemi_baslat, style="success.TButton", width=30).pack(pady=10)

        # Log
        ttk.Label(main_frame, text="İşlem Kayıtları:").pack(anchor=W)
        self.log_area = scrolledtext.ScrolledText(main_frame, height=8, state='disabled', font=("Consolas", 9))
        self.log_area.pack(fill=X, pady=5)

    def log_yaz(self, mesaj):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, mesaj + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')
        self.root.update()

    def word_sec(self):
        dosya = filedialog.askopenfilename(filetypes=[("Word Dosyaları", "*.docx")])
        if dosya: self.word_path.set(dosya)

    def excel_sec(self):
        dosya = filedialog.askopenfilename(filetypes=[("Excel Dosyaları", "*.xlsx")])
        if dosya: self.excel_path.set(dosya)

    def word_degistir(self, doc_obj, eski_metin, yeni_metin):
        for p in doc_obj.paragraphs:
            if eski_metin in p.text:
                p.text = p.text.replace(eski_metin, yeni_metin)
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if eski_metin in p.text:
                            p.text = p.text.replace(eski_metin, yeni_metin)

    def islemi_baslat(self):
        word_dosyasi = self.word_path.get()
        excel_dosyasi = self.excel_path.get()
        konu = self.konu_basligi.get()
        ek_isim = self.dosya_eki.get().strip()
        govde_metni = self.text_area_body.get("1.0", tk.END)

        if not word_dosyasi or not excel_dosyasi:
            messagebox.showerror("Hata", "Lütfen Word ve Excel dosyalarını seçiniz.")
            return

        try:
            df = pd.read_excel(excel_dosyasi)
            df.columns = [c.strip().upper() for c in df.columns]

            if 'ISIM' not in df.columns or 'MAIL' not in df.columns:
                messagebox.showerror("Hata", "Excel dosyasında 'ISIM' ve 'MAIL' sütunları bulunamadı.")
                return
            
            pythoncom.CoInitialize() 
            outlook = win32.Dispatch('outlook.application')

            self.log_yaz("--- İşlem Başlıyor ---")

            for index, row in df.iterrows():
                kisi_ismi = str(row['ISIM'])
                kisi_mail = str(row['MAIL'])
                
                kisi_cc = ""
                if 'CC' in df.columns:
                    val = row['CC']
                    if not pd.isna(val):
                        kisi_cc = str(val)

                self.log_yaz(f"İşleniyor: {kisi_ismi}...")

                # Word İşlemleri
                doc = Document(word_dosyasi)
                self.word_degistir(doc, "{ISIM}", kisi_ismi)

                if ek_isim:
                    pdf_adi = f"{kisi_ismi} {ek_isim}.pdf"
                else:
                    pdf_adi = f"{kisi_ismi}.pdf"

                temp_docx = os.path.abspath(f"temp_{index}.docx")
                temp_pdf = os.path.abspath(pdf_adi)

                doc.save(temp_docx)
                convert(temp_docx, temp_pdf)

                # Mail İşlemleri
                mail = outlook.CreateItem(0)
                mail.To = kisi_mail
                if kisi_cc: mail.CC = kisi_cc
                
                mail.Subject = konu.replace("{ISIM}", kisi_ismi)
                
                mail.Display() 
                imzali_govde = mail.HTMLBody

                guncel_govde = govde_metni.replace("{ISIM}", kisi_ismi)
                guncel_govde_html = guncel_govde.replace("\n", "<br>")

                yeni_html_icerik = f"<div style='font-family: Calibri, Arial; font-size: 11pt;'>{guncel_govde_html}</div><br>" + imzali_govde
                mail.HTMLBody = yeni_html_icerik
                mail.Attachments.Add(temp_pdf)

                mail.Send()
                
                log_mesaji = f"✔ Gönderildi: {kisi_mail}"
                self.log_yaz(log_mesaji)

                try:
                    os.remove(temp_docx)
                    os.remove(temp_pdf) 
                except:
                    pass

            self.log_yaz("--- ✅ TÜM İŞLEMLER TAMAMLANDI ---")
            messagebox.showinfo("Başarılı", "Tüm mailler gönderildi.")

        except Exception as e:
            self.log_yaz(f"❌ HATA: {str(e)}")
            messagebox.showerror("Hata", str(e))

if __name__ == "__main__":
    # TEMA AYARI
    app = ttk.Window(themename="solar") 
    MailGondericiUygulamasi(app)
    app.mainloop()
